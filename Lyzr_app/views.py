# Environment Creation
import json
import os
from datetime import datetime

import chromadb
import requests
from django.conf import settings
from django.core.files.storage import default_storage
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from openai import OpenAI, OpenAIError
from rest_framework.response import Response
from rest_framework.decorators import api_view
from wyge.agents.react_agent import Agent
from wyge.models.openai import ChatOpenAI
from wyge.tools.prebuilt_tools import execute_query

from .database import PostgreSQLDB

db = PostgreSQLDB(dbname='test', user='test_owner', password='tcWI7unQ6REA')


# Create environment
@api_view(['POST'])
def create_environment(request):
    try:
        data = request.data
        name = data.get('name')
        # model_vendor = data.get('model_vendor')
        api_key = data.get('api_key')
        model = data.get('model')
        temperature = data.get('temperature', 0.5)
        # top_p = data.get('top_p', 0.9)

        environment_id = db.create_environment(name, api_key, model, temperature)
        return Response({"environment_id": environment_id}, status=201)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read environment by ID
# Read environment by ID
@api_view(['GET'])
def read_environment(request, environment_id):
    try:
        environment = db.read_environment(environment_id)
        if environment:
            response_data = {

                "features": [],  # Assuming no features are provided, keeping it empty
                "llm_config": {
                    # "provider": environment[2],  # model_vendor
                    "model": environment[3],  # model
                    "config": {
                        "temperature": environment[4],  # temperature
                        # "top_p": environment[6],  # top_p
                    }
                },
                "env": {
                    "Environment_name": environment[1],  # name
                    "OPENAI_API_KEY": environment[2],  # API_KEY
                },

            }

            return Response(response_data, status=200)

        return Response({"error": "Environment not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Update environment by ID
@api_view(['POST'])
def update_environment(request, environment_id):
    try:
        data = request.data
        name = data.get('name')
        # model_vendor = data.get('model_vendor')
        api_key = data.get('api_key')
        model = data.get('model')
        temperature = data.get('temperature')
        # top_p = data.get('top_p')

        updated_rows = db.update_environment(
            environment_id,
            name,
            api_key,
            model,
            temperature

        )

        if updated_rows:
            return Response({"message": f"Environment with ID {environment_id} updated successfully."}, status=200)
        return Response({"message": f"Environment with ID {environment_id} updated successfully."}, status=200)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Delete environment by ID
@api_view(['GET'])
def delete_environment(request, environment_id):
    try:
        deleted_rows = db.delete_environment(environment_id)
        if deleted_rows:
            return Response({"message": f"Environment with ID {environment_id} deleted successfully."}, status=200)
        return Response({"error": f"Environment with ID {environment_id} not found."}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read all environments
@api_view(['GET'])
def read_all_environments(request):
    try:
        environments = db.read_all_environments()
        if environments:
            environment_list = []
            for environment in environments:
                environment_list.append({
                    "id": environment[0],
                    "name": environment[1],
                    # "model_vendor": environment[2],
                    "api_key": environment[2],
                    "model": environment[3],
                    "temperature": environment[4],
                    # "top_p": environment[6],

                })
            return Response(environment_list, status=200)
        return Response({"message": "No environments found."}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


from rest_framework.decorators import api_view
from rest_framework.response import Response
import logging

# Set up logging
logger = logging.getLogger(__name__)


# Create Agent
@api_view(['POST'])
def create_agent(request):
    try:
        data = request.data
        name = data.get('name')
        system_prompt = data.get('system_prompt')
        agent_description = data.get('agent_description')
        tools = data.get('tools')
        upload_attachment = data.get('upload_attachment', False)  # Default value set to False
        env_id = data.get('env_id')
        dynamic_agent_id = data.get('dynamic_agent_id')  # New field

        if not env_id:
            return Response({"error": "Environment ID is required"}, status=400)

        # Create the agent in the database
        agent_id = db.create_agent(
            name, system_prompt, agent_description, tools, upload_attachment, env_id, dynamic_agent_id
        )

        return Response({"agent_id": agent_id}, status=201)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read Agent by ID
@api_view(['GET'])
def read_agent(request, agent_id):
    try:
        agent = db.read_agent(agent_id)
        if agent:
            return Response({
                "id": agent[0],
                "name": agent[1],
                "system_prompt": agent[2],
                "agent_description": agent[3],
                "tools": agent[4],
                "Additional_Features": {
                    "upload_attachment": agent[5],
                },
                "env_id": agent[6],
                "dynamic_agent_id": agent[7]  # Include new field
            }, status=200)
        return Response({"error": "Agent not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Update Agent by ID
@api_view(['POST'])
def update_agent(request, agent_id):
    try:
        data = request.data
        name = data.get('name')
        system_prompt = data.get('system_prompt')
        agent_description = data.get('agent_description')
        tools = data.get('tools')
        upload_attachment = data.get('upload_attachment')
        env_id = data.get('env_id')
        dynamic_agent_id = data.get('dynamic_agent_id')  # Handle new field

        # Update agent in the database
        db.update_agent(
            agent_id, name, system_prompt, agent_description, tools, upload_attachment, env_id, dynamic_agent_id
        )

        return Response({"message": f"Agent with ID {agent_id} updated successfully."}, status=200)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Delete Agent by ID
@api_view(['GET'])
def delete_agent(request, agent_id):
    try:
        db.delete_agent(agent_id)
        return Response({"message": f"Agent with ID {agent_id} deleted successfully."}, status=204)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read All Agents
@api_view(['GET'])
def read_all_agents(request):
    try:
        # Fetch all agents from the database
        agents = db.get_all_agents()

        # Check if any agents are returned
        if not agents:
            return Response({"message": "No agents found"}, status=404)

        # Structure the agents' data for JSON response
        agents_data = [
            {
                "id": agent[0],
                "name": agent[1],
                "system_prompt": agent[2],
                "agent_description": agent[3],
                "tools": agent[4],
                "Additional_Features": {
                    "upload_attachment": agent[5],
                },
                "env_id": agent[6],
                "dynamic_agent_id": agent[7]  # Include new field
            }
            for agent in agents
        ]

        return Response({"agents": agents_data}, status=200)

    except Exception as e:
        # Log the error for further investigation
        logger.error(f"Error fetching agents: {str(e)}")

        # Return a user-friendly error message
        return Response({"error": "An error occurred while fetching agents"}, status=500)


# Creation of the openai environment.
from rest_framework.decorators import api_view
from wyge.agents.prebuilt_agents import ResearchAgent, BlogAgent, LinkedInAgent, ImageGenerationAgent, EmailAgent

from rest_framework.response import Response
from rest_framework import status
import requests
import pandas as pd

from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
import requests


def image_to_base64(image_path):
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode('utf-8')
    except Exception as e:
        return None  # Handle the case where the image path is invalid or the image doesn't exist


# Helper function to delete all images in a directory
def delete_images_in_directory(directory: str) -> None:
    """Delete all image files in the specified directory."""
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
    for filename in os.listdir(directory):
        _, ext = os.path.splitext(filename)
        if ext.lower() in image_extensions:
            os.remove(os.path.join(directory, filename))


# Helper function to get all images in a directory
def get_images_in_directory(directory: str) -> list:
    """Return all image files in the specified directory."""
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
    image_files = []
    for filename in os.listdir(directory):
        _, ext = os.path.splitext(filename)
        if ext.lower() in image_extensions:
            image_files.append(os.path.join(directory, filename))
    return image_files


import re


def differentiate_url(url):
    """
    Differentiates between a general website URL and a YouTube URL.

    Parameters:
    - url (str): The URL to check.

    Returns:
    - str: "YouTube" if the URL is a YouTube link, otherwise "Website".
    """
    youtube_patterns = [
        r"(https?://)?(www\.)?(youtube\.com|youtu\.be)/",  # matches youtube.com or youtu.be
    ]
    for pattern in youtube_patterns:
        if re.match(pattern, url):
            return "YouTube"
    return "Website"


import markdown


def markdown_to_html(md_text):
    html_text = markdown.markdown(md_text)
    return html_text


from dotenv import load_dotenv
import openai

load_dotenv()


@csrf_exempt
def send_email(request):
    """
    API to send the combined blog content and image document via email.
    """
    if request.method == 'POST':
        try:
            # Retrieve data from the POST request
            to_mail = request.POST.get('to_mail')
            if not to_mail:
                return JsonResponse({"error": "Recipient email (to_mail) is required"}, status=400)

            # Path to the combined document (blog content and image)
            combined_doc_file_path = "./blog_post.docx"

            # Validate the attachment
            if not os.path.exists(combined_doc_file_path):
                return JsonResponse({"error": "Combined blog document not found"}, status=400)

            # Send email with the combined document as an attachment
            api_key = os.getenv("OPENAI_API_KEY")
            email_agent = EmailAgent(api_key)
            email_ack = email_agent.send_email(
                to_mail,
                subject="Your Blog Post ",
                body="Here is your generated blog content.",
                attachments=[combined_doc_file_path],
                token_json_file_path="./token.json"
            )

            # Clean up the local combined document after sending the email
            os.remove(combined_doc_file_path)

            return JsonResponse({"ack": email_ack}, status=200)

        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=405)


from rest_framework.response import Response
from rest_framework.decorators import api_view
from rest_framework import status
import os
import base64
import shutil
import tempfile
from django.core.files.storage import default_storage
from django.conf import settings
from wyge.tools.raw_functions import file_to_sql
from .plots import tools as plot_tools
from .generator import generate_synthetic_data, generate_data_from_text, fill_missing_data_in_chunk
from django.http import HttpResponse


# from .utils_for_chat_to_pdf import process_documents_and_store,query_chromadb


@api_view(['POST'])
def run_openai_environment(request):
    try:
        agent_id = request.data.get('agent_id')
        user_prompt = request.data.get('prompt', '')
        url = request.data.get('url', '')
        file = request.FILES.get('file')  # File if attached
        files = request.FILES.getlist('file')

        # Retrieve agent details
        agent = db.read_agent(agent_id)

        # Retrieve the API key from the environment table using env_id
        env_details = db.read_environment(agent[6])  # agent[6] is env_id
        openai_api_key = env_details[2]
        model = env_details[3]  # For chat_app
        temperature = env_details[4]  # For chat_app

        # Process based on input type (prompt, URL, file)
        result = None
        response_data = {}

        # Delete existing image plots (optional)
        delete_images_in_directory(settings.BASE_DIR)

        # Define the list of tool IDs that should be treated as 'blog_post'
        BLOG_TOOL_IDS = ['blog_post', 'mail_blog', 'audio_blog', 'video_blog', 'youtube_blog']

        # Define tool handling conditions
        tool_conditions = {
            'text_to_sql': 'text_to_sql',
            'graph_to_sql': 'graph_to_sql',
            'forecast_to_sql': 'forecast_to_sql'
        }
        # For Text-to-sql ,graph-to-sql and forecast-to-sql
        # Process based on the type of operation in agent[4]
        for tool_key, operation_type in tool_conditions.items():
            if file and tool_key in agent[4]:  # Check if the specific tool is in the agent
                # Call the session-aware function with the current operation_type
                result = handle_excel_file_based_on_type(request, file, openai_api_key, user_prompt, operation_type)

                # Process the response from the function
                if isinstance(result, dict):
                    if 'image_path' in result:
                        response_data["image_base64"] = image_to_base64(result["image_path"])
                    if 'content' in result:
                        response_data["content"] = markdown_to_html(result["content"])
                break  # Exit after the first matching condition is processed

        # Chat_app_with_text
        if user_prompt and 'chat_app_text' in agent[4]:
            audio_response = request.POST.get("audio_response", "false").lower() == "true"
            result = chat_with_openai(openai_api_key, user_prompt, temperature, audio_response)
            # Check if there's an error in the result
            if "error" in result:
                response_data["error"] = result["error"]
            else:
                # Add the text response to the response_data
                response_data["content"] = result["response"]

                # Check if audio_response is True, and include audio in the response_data
                if audio_response:
                    response_data["audio"] = result["audio"]

        # Audio_transcription chat_app
        if 'transcribe_audio' in agent[4]:
            result = transcribe_audio_from_mic(openai_api_key)
            # Check if the transcription was successful
            if "transcription" in result:
                response_data["Transcription"] = result["transcription"]
            else:
                response_data["error"] = result.get("error", "Unknown error occurred")

        # Blog_generation content with urls
        if url and user_prompt:
            print("Getting the url.............")
            url_type = differentiate_url(url)
            print(url_type)
            # Handling based on URL type
            if url_type == "YouTube":
                if any(tool_id in agent[4] for tool_id in BLOG_TOOL_IDS):
                    result = generate_blog_from_yt_url(user_prompt, url, 'blog_post', openai_api_key)
            #     elif 'linkedin_post' in agent[4]:
            #         result = generate_blog_from_yt_url(user_prompt, url, 'linkedin_post', openai_api_key)
            else:  # General Website URL
                if any(tool_id in agent[4] for tool_id in BLOG_TOOL_IDS):
                    result = generate_blog_from_url(user_prompt, url, 'blog_post', openai_api_key)
                # elif 'linkedin_post' in agent[4]:
                #     result = generate_blog_from_url(user_prompt, url, 'linkedin_post', openai_api_key)

            if isinstance(result, dict):
                response_data["content"] = markdown_to_html(result.get("content", ""))
                if "image_path" in result and result["image_path"]:
                    response_data["image_base64"] = image_to_base64(result["image_path"])
                # if "doc_file" in result and result["doc_file"]:
                #     response_data["doc_file"] = result["doc_file"]
                # to_email=request.data.get("email")
                # email_agent = EmailAgent(openai_api_key)
                # ack = email_agent.send_email(
                #     to_email,
                #     'Your content',
                #     'Thank you for using our product.',
                #     result["doc_file"],
                #     token_json_file_path="./token.json"
                # )
                # print(ack)


        # Blog Generation through files
        elif file and user_prompt:
            if any(tool_id in agent[4] for tool_id in BLOG_TOOL_IDS):
                print("function calling here")
                result = generate_blog_from_file(user_prompt, file, 'blog_post', openai_api_key)
                print(result)
            # elif 'linkedin_post' in agent[4]:
            #     result = generate_blog_from_file(user_prompt, file, 'linkedin_post', openai_api_key)

            if isinstance(result, dict):
                response_data["content"] = markdown_to_html(result.get("content", ""))
                if "image_path" in result and result["image_path"]:
                    response_data["image_base64"] = image_to_base64(result["image_path"])
                # if "doc_file" in result and result["doc_file"]:
                #      response_data["doc_file"] =result["doc_file"]
                # to_email=request.data.get("email")
                # email_agent = EmailAgent(openai_api_key)
                # ack = email_agent.send_email(
                #     to_email,
                #     'Your content',
                #     'Thank you for using our product.',
                #     result["doc_file"],
                #     token_json_file_path="./token.json"
                # )
                # print(ack)

        # Synthetic data handling cases(3 cases)
        if user_prompt and 'synthetic_data_new_data' in agent[4]:
            result = handle_synthetic_data_for_new_data(user_prompt, openai_api_key)
            response_data["csv_file"] = result  # Assume result contains CSV file path

        elif file and 'synthetic_data_extended_data' in agent[4]:
            result = handle_synthetic_data_generation(file, user_prompt, openai_api_key)
            response_data["csv_file"] = result

        elif file and 'synthetic_data_missing_data' in agent[4]:
            result = handle_fill_missing_data(file, openai_api_key)
            response_data["csv_file"] = result

        # Chat2pdf function:
        elif files and 'chat_to_pdf' in agent[4]:
            chunk_size = 1500
            result = chat_with_documents(openai_api_key, files, chunk_size, user_prompt)
            response_data["content"] = result["response"]

        # Travel planner
        elif user_prompt and 'travel_planner' in agent[4]:
            weather_api_key = "b307b797aa0caf2cf2c904ae302f7461"
            geolocation_api_key = "2935b537cd024c83a84b7983d7da1ddb"
            result = travel_planning(weather_api_key, geolocation_api_key, openai_api_key, user_prompt)
            response_data["content"] = markdown_to_html(result.get("response", ""))

        # MCQ Generation
        elif user_prompt and 'mcq_generator' in agent[4]:
            result = mcq_generator(openai_api_key, user_prompt)
            # response_data["content"] = result["response"]
            response_data["content"] = markdown_to_html(result.get("response", ""))

        # Interior Design
        elif user_prompt and 'interior_design' in agent[4]:
            print(user_prompt)
            result = interior_designer(user_prompt)
            response_data["content"] = result["response"]

        # #Dynamic_agents
        # elif user_prompt and 'dynamic_agents' in agent[4]:

        # Construct response
        if response_data:
            return Response(response_data, status=status.HTTP_200_OK)
        else:
            return Response({"error": "No valid tool found for the given input."}, status=status.HTTP_400_BAD_REQUEST)

    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)


from docx import Document
from docx.shared import Inches


def save_blog_and_image_to_docx(blog_content, image_path, file_path):
    doc = Document()
    doc.add_heading("Blog Post", level=0)

    # Add the blog content
    doc.add_paragraph(blog_content)

    # Add the image if it exists
    if image_path and os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.0))  # Adjust width as needed

    doc.save(file_path)
    return file_path


def generate_blog_from_url(prompt, url, option, api_key):
    if option == 'blog_post':
        print(datetime.now())
        print("Agent Initialisation will start here........")
        research_agent = ResearchAgent(api_key)
        blog_agent = BlogAgent(api_key)
        img_agent = ImageGenerationAgent(api_key)
        print("crossing all the agent initialisations")
        # Generate content and context
        context = research_agent.research_website(prompt, url)
        blog_content = blog_agent.generate_blog(prompt, context)
        print(datetime.now())
        # img_agent.generate_image(blog_content)
        doc_file, image_path = img_agent.add_to_blog(blog_content)

        # Save blog content and image to a single .docx file
        combined_doc_file_path = "./blog_post.docx"
        save_blog_and_image_to_docx(blog_content, image_path, combined_doc_file_path)

        return {
            "content": blog_content,
            "image_path": image_path,
            "combined_doc_file": combined_doc_file_path
        }
        # return {"content": blog, "image_path": image}
        # elif option == 'linkedin_post':
        #     linkedin_agent = LinkedInAgent(api_key)
        #     research_agent = ResearchAgent(api_key)
        #     img_agent = ImageGenerationAgent(api_key)
        #     context = research_agent.research_website(prompt, url)
        #     content = linkedin_agent.generate_linkedin_post(context)
        #     img_agent.generate_image(content)
        #     doc_file, image_path = img_agent.add_to_blog(content)
        #     return {"content": content, "image_path": image_path}
    # except Exception as e:
    #     return {"error": str(e)}


# Generate content from URL (for blog or LinkedIn post)
def generate_blog_from_yt_url(prompt, url, option, api_key):
    try:
        if option == 'blog_post':
            research_agent = ResearchAgent(api_key)
            blog_agent = BlogAgent(api_key)
            img_agent = ImageGenerationAgent(api_key)
            print(url, datetime.now())

            blog_content = research_agent.extract_transcript_from_yt_video(url)
            print("Blog content is.............")
            context = blog_agent.generate_blog(prompt, blog_content)
            # img_agent.generate_image(context)
            doc_file, image_path = img_agent.add_to_blog(context)

            print("result", datetime.now())
            # Save blog content and image to a single .docx file
            combined_doc_file_path = "./blog_post.docx"
            print(image_path)
            save_blog_and_image_to_docx(context, image_path, combined_doc_file_path)
            print(image_path)
            return {
                "content": blog_content,
                "image_path": image_path,
                "combined_doc_file": combined_doc_file_path
            }

            # return {"content": blog, "image_path": image}
        # elif option == 'linkedin_post':
        #     linkedin_agent = LinkedInAgent(api_key)
        #     yt_agent = YTBlogAgent(api_key)
        #     context = yt_agent.extract_transcript(url)
        #     content, image_path = linkedin_agent.generate_linkedin_post(context)
        #     return {"content": content, "image_path": image_path}
    except Exception as e:
        return {"error": str(e)}


# Generate content from file (for blog or LinkedIn post)
def generate_blog_from_file(prompt, file, option, api_key):
    try:
        print("*******")
        file_path = save_file(file)
        print(file_path)
        if option == 'blog_post':
            print("started")
            research_agent = ResearchAgent(api_key)
            blog_agent = BlogAgent(api_key)
            img_agent = ImageGenerationAgent(api_key)
            print("getting started")

            blog_content = research_agent.extract_text_from_audio_or_video(file_path)
            context = blog_agent.generate_blog(prompt, blog_content)
            # img_agent.generate_image(context)
            doc_file, image_path = img_agent.add_to_blog(context)

            # Save blog content and image to a single .docx file
            combined_doc_file_path = "./blog_post.docx"
            save_blog_and_image_to_docx(context, image_path, combined_doc_file_path)

            return {
                "content": blog_content,
                "image_path": image_path,
                "combined_doc_file": combined_doc_file_path
            }
            # return {"content": blog, "image_path": image}

        # elif option == 'linkedin_post':
        #     linkedin_agent = LinkedInAgent(api_key)
        #     va_agent = VideoAudioBlogAgent(api_key)
        #     context = va_agent.extract_text(file_path)
        #     content, image_path = linkedin_agent.generate_linkedin_post(context)
        #     return {"content": content, "image_path": image_path}

    except Exception as e:
        return {"error": str(e)}


def extract_num_rows_from_prompt(user_prompt):
    """
    Extracts the number of rows or records from the user's prompt.
    Supports prompts like:
      - "Generate 100 rows of data"
      - "Generate the 50 records of data"
    """
    match = re.search(r'(\d+)\s+(rows|records)', user_prompt, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def extract_columns_from_prompt(user_prompt):
    """
    Extracts the field names (column names) from the user's prompt.
    Supports prompts like:
      - "field names: S.no, Name, address, First_name"
      - "fields: S.no, Name, address, First_name"
      - "columns: S.no, Name, address, First_name"
      - "field_names: S.no, Name, address, First_name"
      - "column_names: S.no, Name, address, First_name"

    Converts column names to snake_case, removes spaces and special characters.
    Example:
      - "S.no, Name, address, First_name"
      -> ['s_no', 'name', 'address', 'first_name']
    """
    # Look for all possible field identifier formats followed by the column names
    match = re.search(r'(field names|column names|fields|columns|field_names|column_names):\s*([a-zA-Z0-9_,\s\.]+)',
                      user_prompt, re.IGNORECASE)

    if match:
        # Extract the part containing column names
        raw_columns = match.group(2).split(',')
    else:
        return []

    # Format each column name (remove spaces, convert to snake_case, lowercase)
    formatted_columns = [
        re.sub(r'[^a-zA-Z0-9]', '_', col.strip()).lower()
        for col in raw_columns
    ]

    # Remove empty column names and ensure no duplicates
    formatted_columns = list(filter(bool, formatted_columns))
    return list(dict.fromkeys(formatted_columns))  # Remove duplicates


def handle_synthetic_data_for_new_data(user_prompt, openai_api_key):
    """
    Function to handle synthetic data generation.

    Parameters:
    - uploaded_file: The empty Excel or CSV file containing column names
    - user_prompt: The user's prompt, which should contain the number of rows
    - openai_api_key: The OpenAI API key to be used for generating synthetic data

    Returns:
    - A CSV string with generated synthetic data or an error message
    """
    try:
        # # Determine file type and extract column names
        # file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        # print(file_extension)
        # if file_extension == ".xlsx":
        #     df = pd.read_excel(uploaded_file)
        # elif file_extension == ".csv":
        #     df = pd.read_csv(uploaded_file)
        # else:
        #     return {"error": "Unsupported file format. Please upload an Excel or CSV file."}
        #
        # column_names = df.columns.tolist()
        #
        # # Check if the necessary information is provided
        # if not user_prompt or not column_names:
        #     return {"error": "Missing user prompt or column names"}

        # Validate user prompt
        if not user_prompt or not openai_api_key:
            return JsonResponse({"error": "Missing required parameters: user_prompt or OpenAI API key"}, status=400)

        # Extract number of rows from the prompt
        num_rows = extract_num_rows_from_prompt(user_prompt)
        if num_rows is None:
            return JsonResponse({"error": "Number of rows or records not found in the prompt."}, status=400)

        # Extract column names from the prompt
        column_names = extract_columns_from_prompt(user_prompt)
        print(column_names)
        if not column_names:
            return JsonResponse({"error": "No field names found in the prompt."}, status=400)

        # Generate synthetic data using the column names and the number of rows
        generated_df = generate_data_from_text(openai_api_key, user_prompt, column_names, num_rows=num_rows)

        # Convert the generated DataFrame to CSV format
        combined_csv = generated_df.to_csv(index=False)

        return {
            "data": combined_csv
        }

    except Exception as e:
        return {"error": str(e)}


def handle_synthetic_data_generation(file, user_prompt, openai_api_key):
    """
    Function to handle synthetic data generation and merging with original data.

    Parameters:
    - file: The uploaded Excel or CSV file containing original data
    - user_prompt: The user's prompt, which should contain the number of rows
    - openai_api_key: The OpenAI API key to be used for generating synthetic data

    Returns:
    - A dictionary with message and combined CSV data, or an error message
    """
    try:
        # Determine file type and read into a DataFrame
        file_extension = os.path.splitext(file.name)[1].lower()
        print(file_extension)
        if file_extension == ".xlsx":
            original_df = pd.read_excel(file)
        elif file_extension == ".csv":
            original_df = pd.read_csv(file)
            print(original_df)
        else:
            return {"error": "Unsupported file format. Please upload an Excel or CSV file."}

        # Check if the DataFrame has more than 300 rows
        if len(original_df) > 300:
            original_df = original_df.head(300)  # Take the first 300 rows

        # Create a temporary file for the data
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
            temp_file_name = temp_file.name
            # Save the truncated or original data to a temporary location
            if file_extension == ".xlsx":
                original_df.to_excel(temp_file_name, index=False)
            elif file_extension == ".csv":
                original_df.to_csv(temp_file_name, index=False)

        # Extract the number of rows from the prompt
        num_rows = extract_num_rows_from_prompt(user_prompt)
        if num_rows is None:
            return {"error": "Number of rows not found in the prompt"}

        # Generate synthetic data using the temporary file path
        generated_df = generate_synthetic_data(openai_api_key, temp_file_name, num_rows)

        # Combine the original and synthetic data
        combined_df = pd.concat([original_df, generated_df], ignore_index=True)

        # Convert to CSV for download
        combined_csv = combined_df.to_csv(index=False)

        return {
            "data": combined_csv
        }

    except Exception as e:
        return {"error": str(e)}


def handle_fill_missing_data(file, openai_api_key):
    """
    Function to handle filling missing data in chunks and returning as CSV.

    Parameters:
    - file: The uploaded Excel or CSV file containing original data
    - openai_api_key: The OpenAI API key to be used for generating missing data

    Returns:
    - A dictionary with a message and combined CSV data, or an error message
    """
    try:
        # Determine file type and read into a DataFrame
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension == ".xlsx":
            original_df = pd.read_excel(file)
        elif file_extension == ".csv":
            original_df = pd.read_csv(file)
        else:
            return {"error": "Unsupported file format. Please upload an Excel or CSV file."}

        # Check if the DataFrame has more than 300 rows
        if len(original_df) > 300:
            original_df = original_df.head(300)  # Limit to the first 300 rows

        # Create a temporary file for the data
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
            temp_file_name = temp_file.name
            # Save the truncated or original data to a temporary location
            if file_extension == ".xlsx":
                original_df.to_excel(temp_file_name, index=False)
            elif file_extension == ".csv":
                original_df.to_csv(temp_file_name, index=False)

        # Fill missing data using the temporary file path
        filled_df = fill_missing_data_in_chunk(openai_api_key, temp_file_name)

        # Convert the filled DataFrame to CSV format for output
        combined_csv = filled_df.to_csv(index=False)

        return {
            "data": combined_csv
        }

    except Exception as e:
        return {"error": str(e)}


# Text to sql function
import os
from django.core.files.storage import default_storage
from django.http import HttpResponse
from wyge.models.openai import ChatOpenAI
from wyge.agents.react_agent import Agent
from wyge.tools.prebuilt_tools import execute_query, execute_code, install_library
from wyge.tools.raw_functions import file_to_sql, get_metadata
from .system_prompt1 import reAct_prompt
from .system_prompt2 import plot_prompt
from .system_prompt3 import forecasting_prompt

# Shared database credentials
USER = 'test_owner'
PASSWORD = 'tcWI7unQ6REA'
HOST = 'ep-yellow-recipe-a5fny139.us-east-2.aws.neon.tech:5432'
DATABASE = 'test'


def handle_excel_file_based_on_type(request, file, openai_api_key, user_prompt, operation_type):
    """
    Processes an uploaded Excel file by storing it in the database, converting it to SQL,
    and generating output based on the operation type (text, graph, forecast).

    Parameters:
    - file: File object of the uploaded Excel file
    - openai_api_key: OpenAI API key for model access
    - user_prompt: Query to be processed by the AI agent
    - operation_type: Specifies the operation type ('text_to_sql', 'graph_to_sql', 'forecast_to_sql')

    Returns:
    - Dictionary with content and optionally an image path
    """
    table_name = file.name.split('.')[0]
    # text-to-sql session starts here
    print(datetime.now())
    if 'processed_tables' not in request.session:
        request.session['processed_tables'] = []

        # If the table hasn't been processed, save file and convert to SQL
    if table_name not in request.session['processed_tables']:
        # Save the uploaded file to a storage path
        file_path = os.path.join(settings.MEDIA_ROOT, file.name)
        with default_storage.open(file_path, 'wb+') as f:
            for chunk in file.chunks():
                f.write(chunk)

        # Convert the Excel file to SQL table (runs only once per table)
        file_to_sql(file_path, table_name, USER, PASSWORD, HOST, DATABASE)
        request.session['processed_tables'].append(table_name)

    print("Current session contents:")
    for key, value in request.session.items():
        print(f"{key}: {value}")

    # Select tools and prompt based on operation type
    if operation_type == 'text_to_sql':
        print("Going in to the database............")
        tools = [execute_query()]
        print("Fetching from the database.........")
        prompt = reAct_prompt
    elif operation_type == 'graph_to_sql':
        tools = [execute_query(), execute_code(), install_library()]
        prompt = plot_prompt
    elif operation_type == 'forecast_to_sql':
        tools = [execute_query(), execute_code(), install_library()]
        prompt = forecasting_prompt
    else:
        return JsonResponse({"error": "Invalid operation type specified."}, status=400)

    print("--------------------------------------------------")

    llm = ChatOpenAI(memory=True, tools=tools, api_key=openai_api_key)
    if 'llm_mem' in request.session:
        llm.chat_memory.memory = request.session['llm_mem']

    agent = Agent(llm, react_prompt=prompt)

    # Retrieve or generate metadata
    if 'llm_metadata' not in request.session:
        metadata = get_metadata(HOST, USER, PASSWORD, DATABASE, request.session['processed_tables'])
        request.session['llm_metadata'] = metadata
        print("Initialized new metadata.")
    else:
        metadata = request.session['llm_metadata']
        print("Reusing existing metadata.")

    # Prepare command for agent execution
    print("Preparing command for the agent:")
    print(tools)
    command = f"""
        Answer the user query from the database below,also use the provided tools.
        user = '{USER}'
        password = '{PASSWORD}'
        host = '{HOST}'
        database = '{DATABASE}'
        tables related to user are: {request.session['processed_tables']}
        Metadata of the tables: {metadata}
        User query: {user_prompt}
    """
    print("Command:", command)

    # Execute the command with the agent
    try:
        print("----------------------------------------------------------")
        print(datetime.now())
        print(agent.llm.tools)
        print(agent.react_prompt)
        response = agent(command)
        print("----------------------------------------------------------")
        print(datetime.now())
        print("Raw agent response:", response)
        response = response.split('Final Answer:')[-1]
    except Exception as e:
        print("Error executing agent command:", str(e))
        return JsonResponse({"error": "Failed to execute command."}, status=500)

    # Fetch any generated images
    images = get_images_in_directory(settings.BASE_DIR)

    # Prepare result
    result = {"content": response}
    if images:
        result["image_path"] = images[0]  # Only add image_path if images are available

    # Update the session with the new memory state
    request.session['llm_memory'] = llm.memory  # Store only memory, not the whole instance
    request.session['llm_mem'] = remove_tool_entries(remove_non_serializable(llm.chat_memory.get_memory()))
    print(request.session['llm_mem'])
    print('Updated LLM memory in session:')
    print(request.session['llm_memory'])

    return result


def remove_tool_entries(conversation):
    return [entry for entry in conversation if entry['role'] != 'tool']


def remove_non_serializable(data_list):
    def is_serializable(item):
        try:
            json.dumps(item)
            return True
        except (TypeError, OverflowError):
            return False

    def filter_serializable(item):
        if isinstance(item, dict):
            # Filter out non-serializable values
            return {k: filter_serializable(v) for k, v in item.items() if is_serializable(v)}
        elif isinstance(item, list):
            # Filter out non-serializable items
            return [filter_serializable(i) for i in item if is_serializable(i)]
        return item if is_serializable(item) else None

    # Create a filtered list and remove any None or empty dicts
    filtered_result = [filter_serializable(item) for item in data_list]

    # Remove None values and empty dictionaries from the final result
    return [item for item in filtered_result if
            item and (isinstance(item, dict) and item) or (isinstance(item, list) and item)]
    # Helper function to save uploaded file


# Chat_app_with text
def chat_with_openai(api_key, user_prompt, temperature, audio_response):
    try:
        # Initialize the language model
        print(audio_response)
        llm = ChatOpenAI(api_key=api_key, temperature=temperature)
        response_text = llm.run(user_prompt)
    except OpenAIError as e:
        return {"error": f"OpenAI API error: {str(e)}"}

    # Generate audio response if requested
    if audio_response:
        try:
            tts_client = OpenAI(api_key=api_key)
            tts_response = tts_client.audio.speech.create(
                model="tts-1",
                voice="alloy",
                input=response_text,
            )
            audio_base64 = base64.b64encode(tts_response.content).decode("utf-8")
            return {"response": response_text, "audio": audio_base64}
        except OpenAIError as e:
            return {"error": f"Text-to-speech error: {str(e)}"}

    # Return only the text response if audio is not requested
    print("Response is..................")
    print(response_text)
    return {"response": response_text}


import speech_recognition as sr


# Transcribe audio for chat_app using microphone input
def transcribe_audio_from_mic(api_key):
    recognizer = sr.Recognizer()

    # Use microphone as the audio source
    with sr.Microphone() as source:
        print("Please speak into the microphone...")
        audio_data = recognizer.listen(source)  # Capture audio from the microphone

    try:
        # Convert the audio data to text using Whisper API
        openai_client = OpenAI(api_key=api_key)
        audio_bytes = audio_data.get_wav_data()  # Get audio data in WAV format
        file_tuple = ("audio.wav", audio_bytes)  # Prepare as a tuple for OpenAI API

        # Transcribe using Whisper model
        transcript = openai_client.audio.transcriptions.create(
            model="whisper-1",
            file=file_tuple
        )

        print("Transcript printed:", transcript)

        # Access the transcription text if available
        if hasattr(transcript, 'text'):
            return {"transcription": transcript.text}
        else:
            return {"error": "Transcription object does not have 'text' attribute"}

    except OpenAIError as e:
        return {"error": f"Transcription error: {str(e)}"}

    except sr.UnknownValueError:
        return {"error": "Could not understand audio"}

    except sr.RequestError as e:
        return {"error": f"Speech Recognition error: {str(e)}"}


# Chat2pdf function code
from qdrant_client import QdrantClient
from .utils import process_documents_and_store, query_qdrant
from qdrant_client.http import models


def chat_with_documents(api_key, files, chunk_size, user_prompt):
    # Initialize Qdrant client and OpenAI model with necessary configurations
    processed_files = []
    client = QdrantClient(
        url="https://e3e1ed24-80eb-4e12-98a3-22020369714b.us-east4-0.gcp.cloud.qdrant.io:6333",
        api_key="VvJp11LzgE25Zmj1SMRcxsTVRstzBCwF1UYvLMPUYWww4hO9WGKegg"
    )
    llm = ChatOpenAI(memory=True, api_key=api_key)

    # Ensure collection is created
    collection_name = "document_embeddings"
    all_collections = [col.name for col in client.get_collections().collections]
    if collection_name in all_collections:
        client.delete_collection(collection_name=collection_name)
    client.create_collection(
        collection_name=collection_name,
        vectors_config=models.VectorParams(size=384, distance=models.Distance.COSINE)
    )
    for file in files:
        file_path = save_file(file)
        with open(file_path, 'rb') as f:
            # Process and store document embeddings in Qdrant
            process_documents_and_store(client, collection_name, [f], chunk_size)
        os.remove(file_path)  # Clean up uploaded file
        processed_files.append(f"{file.name} processed successfully.")
    print("Files processed successfully")

    # Retrieve relevant document sections from Qdrant
    results = query_qdrant(client, collection_name, user_prompt)
    context_docs = "\n".join([doc[1]["text"] for doc in results])

    prompt = f"""Use the provided context to answer the user questions. The entire context may not be related to user question, so answer wisely from the context.
    \nIf the answer is not available in the context, please respond with "I couldn't find relevant information about that in the provided documents."

    \ncontext:
    \n{context_docs}

    \nuser query: {user_prompt}
    """

    # Get the response from the language model
    response = llm.run(prompt)
    return {"response": response}


# Travel Planning
from .traveller_planer import TravelPlannerAgent


def travel_planning(weather_api_key, geolocation_api_key, openai_api_key, user_prompt):
    # Initialize the travel planner agent with API keys
    travel_agent = TravelPlannerAgent(
        weather_api_key=weather_api_key,
        geolocation_api_key=geolocation_api_key,
        openai_api_key=openai_api_key
    )
    travel_plan = travel_agent.generate_travel_plan(user_prompt)
    return {"response": travel_plan}


# MCQ generation
from .mcq import MCQGeneratorAgent


def mcq_generator(openai_api_key, user_prompt):
    mcq_agent = MCQGeneratorAgent(openai_api_key)
    # Generate MCQs based on the prompt
    mcq_set = mcq_agent.generate_mcq_set(user_prompt)
    return {"response": mcq_set}


def save_file(file):
    directory = 'temp'
    if not os.path.exists(directory):
        os.makedirs(directory)

    file_path = os.path.join(directory, file.name)
    with open(file_path, "wb") as f:
        f.write(file.read())
    return file_path


# Interior design agent
from .interior import InteriorDesignAgent


def interior_designer(user_prompt):
    agent = InteriorDesignAgent()
    image_url = agent.generate_design(user_prompt)
    print("Generated Image URL:", image_url)
    return {"response": image_url}


# -----------------------------------------------------------------------------------------------------------

# Dynamic agen creation code

# Create Agent
@api_view(['POST'])
def create_dynamic_agent(request):
    """
    API to create an agent dynamically based on the dynamic_ai_agents table.
    """
    try:
        # Extract agent details from the request
        data = request.data
        name = data.get('agent_name')
        agent_goal = data.get('agent_goal')  # Change from 'agent_role' to 'agent_goal'
        agent_description = data.get('agent_description')
        agent_instruction = data.get('agent_instruction')

        # Ensure required fields are provided
        if not name or not agent_goal or not agent_description:
            return Response({"error": "Agent name, goal, and description are required."}, status=400)

        # Step 1: Create the agent in the database
        agent_id = db.create_dynamic_agent(name, agent_goal, agent_description, agent_instruction)
        if not agent_id:
            return Response({"error": "Failed to create agent in the database."}, status=500)

        # Retrieve the created agent details
        agent = db.read_dynamic_agent(agent_id)
        if not agent:
            return Response({"error": "Agent details not found or incomplete."}, status=500)

        # Step 2: Construct agent details for the response
        agent_details = {
            "agent_id": agent_id,
            "agent_name": agent[1],
            "agent_goal": agent[2],
            "agent_description": agent[3],
            "agent_instruction": agent[4]
        }

        # Step 3: Return success response with agent details
        return Response({
            "message": "Agent created successfully.",
            "agent_details": agent_details
        }, status=201)

    except Exception as e:
        logger.error(f"Error creating agent: {str(e)}")  # Add detailed logging
        return Response({"error": f"Internal error: {str(e)}"}, status=400)


# Read Agent by ID
@api_view(['GET'])
def read_dynamic_agent(request, agent_id):
    try:
        agent = db.read_dynamic_agent(agent_id)
        if agent:
            return Response({
                "id": agent[0],
                "agent_name": agent[1],
                "agent_goal": agent[2],
                "agent_description": agent[3],
                "agent_instruction": agent[4]

            }, status=200)
        return Response({"error": "Agent not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Update Agent by ID
@api_view(['POST'])
def update_dynamic_agent(request, agent_id):
    try:
        data = request.data
        name = data.get('agent_name')
        agent_goal = data.get('agent_goal')
        agent_description = data.get('agent_description')
        agent_instruction = data.get('agent_instruction')

        # Update agent in the database
        db.update_dynamic_agent(agent_id, name, agent_goal, agent_description, agent_instruction)

        return Response({"message": f"Agent with ID {agent_id} updated successfully."}, status=200)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Delete Agent by ID
@api_view(['GET'])
def delete_dynamic_agent(request, agent_id):
    try:
        db.delete_dynamic_agent(agent_id)
        return Response({"message": f"Agent with ID {agent_id} deleted successfully."}, status=204)
    except Exception as e:
        return Response({"error": str(e)}, status=400)


# Read all Agents
from rest_framework.decorators import api_view
from rest_framework.response import Response
import logging


@api_view(['GET'])
def read_all_dynamic_agents(request):
    try:
        # Fetch all agents from the database
        agents = db.get_all_dynamic_agents()

        # Check if any agents are returned
        if not agents:
            return Response({"message": "No agents found"}, status=404)

        # Structure the agents' data for JSON response
        agents_data = [
            {
                "id": agent[0],
                "agent_name": agent[1],
                "agent_role": agent[2],
                "agent_description": agent[3],
                "agent_instruction": agent[4]
            }
            for agent in agents
        ]

        return Response({"agents": agents_data}, status=200)

    except Exception as e:
        # Log the error for further investigation
        logger.error(f"Error fetching agents: {str(e)}")

        # Return a user-friendly error message
        return Response({"error": "An error occurred while fetching agents"}, status=500)


# Main API to create an agent(dynamically)
@api_view(['POST'])
def create_openai_environment_api(request):
    """
    API to create an OpenAI environment dynamically based on agent_name, agent_goal, and agent_description.
    """
    try:
        agent_id = request.data.get('agent_id')

        if not agent_id:
            return Response({"error": "Agent ID is required"}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve the agent details from the database
        agent = db.read_dynamic_agent(agent_id)
        if not agent:
            return Response({"error": "Agent not found"}, status=status.HTTP_404_NOT_FOUND)

        # Extract agent details
        agent_details = {
            "name": agent[1],
            "agent_goal": agent[2],
            "agent_description": agent[3],
            "ext_tools": agent[4],
            "env_id": agent[5]
        }

        # Retrieve API key from the environment table
        env_details = db.read_environment(agent_details['env_id'])
        if not env_details or not env_details[2]:
            return Response({"error": "API key not found in environment table"},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        openai_api_key = env_details[2]

        # Dynamically construct the system prompt based on agent details
        system_prompt = generate_system_prompt(agent_details)

        # Add the system prompt to agent details for payload
        agent_details['system_prompt'] = system_prompt

        # Create the OpenAI environment
        environment_response = create_openai_environment(agent_details, openai_api_key)

        if environment_response.get("success"):
            return Response({"message": "OpenAI environment created successfully.",
                             "details": environment_response},
                            status=status.HTTP_201_CREATED)
        else:
            return Response({"error": "Failed to create OpenAI environment.",
                             "details": environment_response},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)


def generate_system_prompt(agent_details):
    return f"""
    You are {agent_details['name']}, an AI agent designed to achieve the following goal:
    {agent_details['agent_goal']}.

    Your description: {agent_details['agent_description']}.
    
    You have to follow the instructions:{agent_details['agent_instructions']}

    Behave in a manner consistent with this goal and description. Use external tools if required 
    (e.g., {', '.join(agent_details['tools']) if agent_details['tools'] else 'none'}). 
    Provide accurate, helpful, and concise responses to user queries.
    """


def create_openai_environment(agent_details, openai_api_key):
    try:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "gpt-4o-mini",  # Adjust the model based on your requirements
            "messages": [
                {"role": "system", "content": agent_details['system_prompt']},
                {"role": "user", "content": agent_details['agent_description']}
            ],
            "max_tokens": 1500
        }

        response = requests.post(url, headers=headers, json=payload)

        if response.status_code == 200:
            return {"success": True, "response": response.json()}
        else:
            return {"success": False, "error": response.text}

    except Exception as e:
        return {"success": False, "error": str(e)}


# API: Run OpenAI Environment
# @api_view(['POST'])
# def run_agent_environment(request):
#     """
#     API to interact with a dynamically created OpenAI environment using user-specific queries.
#     """
#     try:
#         print("Received request to run_agent_environment")
#
#         # Extract the agent ID and user query from the request
#         data = request.data
#         print("Request data:", data)
#
#         agent_id = data.get('agent_id')
#         user_query = data.get('query')
#
#         if not agent_id or not user_query:
#             print("Missing agent_id or query in request")
#             return Response({"error": "Agent ID and query are required"}, status=400)
#
#         print("Fetching agent details for agent_id:", agent_id)
#         # Retrieve agent details from the database
#         agent = db.read_agent(agent_id)
#         if not agent:
#             print("Agent not found for agent_id:", agent_id)
#             return Response({"error": "Agent not found"}, status=404)
#
#         dyn_agent=agent[7]
#         dynamic_agent=db.read_dynamic_agent(dyn_agent)
#
#         # Extract agent details
#         agent_details = {
#             "name": dynamic_agent[1],
#             "agent_goal": dynamic_agent[2],
#             "agent_description": dynamic_agent[3],
#             "agent_instructions":dynamic_agent[4],
#             "tools":agent[4]
#         }
#         print("Agent details:", agent_details)
#
#         # Retrieve API key from the environment table
#         print("Fetching environment details for env_id:", agent[6])
#         env_details = db.read_environment(agent[6])
#         if not env_details or not env_details[2]:
#             print("API key not found in environment table for env_id:", agent_details['env_id'])
#             return Response({"error": "API key not found in environment table"}, status=500)
#
#         openai_api_key = env_details[2]
#         print("OpenAI API key retrieved successfully")
#
#         # Prepare payload for OpenAI API request
#         url = "https://api.openai.com/v1/chat/completions"
#         headers = {
#             "Authorization": f"Bearer {openai_api_key}",
#             "Content-Type": "application/json"
#         }
#         payload = {
#             "model": "gpt-4o-mini",  # Adjust the model based on your requirements
#             "messages": [
#                 {"role": "system", "content": generate_system_prompt(agent_details)},
#                 {"role": "user", "content": user_query}
#             ],
#             "max_tokens": 1500
#         }
#         print("Payload prepared for OpenAI API request:", payload)
#
#         # Send the query to the OpenAI environment
#         print("Sending request to OpenAI API...")
#         response = requests.post(url, headers=headers, json=payload)
#         print("OpenAI API response status code:", response.status_code)
#
#         if response.status_code == 200:
#             openai_response = response.json()
#             print("OpenAI API response:", openai_response)
#
#             # Extract the 'choices' key from the OpenAI response to get the content
#             chat_content = openai_response.get("choices", [{}])[0].get("message", {}).get("content", "")
#             print("Extracted chat content:", chat_content)
#
#             return Response({
#                 "success": True,
#                 "response": markdown_to_html(chat_content)
#             }, status=200)
#         else:
#             print("Error response from OpenAI API:", response.text)
#             return Response({
#                 "success": False,
#                 "error": response.text
#             }, status=500)
#
#     except Exception as e:
#         print("Exception occurred:", str(e))
#         return Response({"error": str(e)}, status=400)

from wyge.agents.prebuilt_agents import GoogleDriveAgent, EmailAgent


@api_view(['POST'])
def run_agent_environment(request):
    """
    API to interact with a dynamically created OpenAI environment using user-specific queries.
    """
    try:
        print("Received request to run_agent_environment")

        # Extract the agent ID and user query from the request
        data = request.data
        print("Request data:", data)

        agent_id = data.get('agent_id')
        user_query = data.get('query')
        recipient_email = data.get('email')  # Email comes from the frontend
        authorisation_code = data.get('auth_code')

        if not agent_id or not user_query:
            print("Missing agent_id or query in request")
            return Response({"error": "Agent ID and query are required"}, status=400)

        print("Fetching agent details for agent_id:", agent_id)
        # Retrieve agent details from the database
        agent = db.read_agent(agent_id)
        if not agent:
            print("Agent not found for agent_id:", agent_id)
            return Response({"error": "Agent not found"}, status=404)

        dyn_agent = agent[7]
        dynamic_agent = db.read_dynamic_agent(dyn_agent)

        # Extract agent details
        agent_details = {
            "name": dynamic_agent[1],
            "agent_goal": dynamic_agent[2],
            "agent_description": dynamic_agent[3],
            "agent_instructions": dynamic_agent[4],
            "tools": agent[4]
        }
        print("Agent details:", agent_details)

        # Retrieve API key from the environment table
        print("Fetching environment details for env_id:", agent[6])
        env_details = db.read_environment(agent[6])
        if not env_details or not env_details[2]:
            print("API key not found in environment table for env_id:", agent_details['env_id'])
            return Response({"error": "API key not found in environment table"}, status=500)

        openai_api_key = env_details[2]
        print("OpenAI API key retrieved successfully")

        # Prepare payload for OpenAI API request
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "gpt-4o-mini",  # Adjust the model based on your requirements
            "messages": [
                {"role": "system", "content": generate_system_prompt(agent_details)},
                {"role": "user", "content": user_query}
            ],
            "max_tokens": 1500
        }
        print("Payload prepared for OpenAI API request:", payload)

        # Send the query to the OpenAI environment
        print("Sending request to OpenAI API...")
        response = requests.post(url, headers=headers, json=payload)
        print("OpenAI API response status code:", response.status_code)

        if response.status_code == 200:
            openai_response = response.json()
            print("OpenAI API response:", openai_response)

            # Extract the 'choices' key from the OpenAI response to get the content
            chat_content = openai_response.get("choices", [{}])[0].get("message", {}).get("content", "")
            print("Extracted chat content:", chat_content)
            markdown_content = markdown_to_html(chat_content)

            # Create a .docx file for the response
            file_name = "response.docx"
            file_path = f"./{file_name}"
            try:
                from docx import Document
                document = Document()
                document.add_heading("Response Content", level=1)
                document.add_paragraph(chat_content)
                document.save(file_path)
                print(f"{file_name} created successfully.")
            except Exception as e:
                return Response({"error": f"Failed to create .docx file: {str(e)}"}, status=500)

            # Handle email or Google Drive based on the provided ID
            if "send_mail" in agent[4]:
                # Use EmailAgent to send the .docx file as an attachment
                print("Using EmailAgent to send email...")
                email_agent = EmailAgent(api_key=openai_api_key)
                email_ack = email_agent.send_email(
                    to_mail=recipient_email,
                    subject=f"Response from {agent_details['name']}",
                    body="Please find the response attached.",
                    attachments=[file_path],
                    credentials_json_file_path="credentials.json",
                    token_json_file_path="token.json"
                )
                print("Email sent acknowledgment:", email_ack)

            # elif "send_to_drive" in agent[4]:
            #     # Use GoogleDriveAgent to upload the .docx file
            #     print("Using GoogleDriveAgent to upload file...")
            #     drive_agent = GoogleDriveAgent(api_key=openai_api_key)
            #     drive_ack = drive_agent.upload(
            #         file_path=file_path,
            #         file_name=file_name,
            #         parent_folder_id=parent_folder_id,
            #         service_account="drive_api_access_key.json"
            #     )
            #     print("Google Drive upload acknowledgment:", drive_ack)

            if "send_to_drive" in agent[4]:
                print("Handling Google Drive upload using user's authorization code...")
                drive_file_id = upload_to_drive(file_path, authorisation_code)
                print("sent successfully")
                if not drive_file_id:
                    return Response({"error": "Failed to upload file to Google Drive"}, status=500)

            # Cleanup the local .docx file
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Local file {file_name} deleted successfully.")

            return Response({
                "success": True,
                "message": "Response processed successfully.",
                "content": markdown_content
            }, status=200)
        else:
            print("Error response from OpenAI API:", response.text)
            return Response({
                "success": False,
                "error": response.text
            }, status=500)

    except Exception as e:
        print("Exception occurred:", str(e))
        return Response({"error": str(e)}, status=400)

def upload_to_drive(file_path, authorization_code):
    try:
        import os
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaFileUpload

        CLIENT_SECRETS_FILE = "google_credentials.json"
        REDIRECT_URI = "http://localhost:8000/oauth2callback/"

        # Debug: Print the file path and authorization code
        print(f"File path: {file_path}")
        print(f"Authorization code: {authorization_code}")

        if not os.path.exists(CLIENT_SECRETS_FILE):
            raise FileNotFoundError(f"Client secrets file not found at {CLIENT_SECRETS_FILE}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File to upload not found at {file_path}")

        # Exchange the authorization code for access tokens
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=["https://www.googleapis.com/auth/drive.file"],
            redirect_uri=REDIRECT_URI
        )
        flow.fetch_token(code=authorization_code)

        # Debug: Print the credentials information
        credentials = flow.credentials
        print(f"Access token: {credentials.token}")
        print(f"Refresh token: {credentials.refresh_token}")

        drive_service = build('drive', 'v3', credentials=credentials)

        # Debug: Print file metadata being used for upload
        file_metadata = {
            'name': os.path.basename(file_path),
        }
        print(f"File metadata: {file_metadata}")

        media = MediaFileUpload(
            file_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        print(f"Media upload object created: {media}")

        # Attempt to upload the file
        drive_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        # Debug: Print the response from the API
        print(f"Drive API response: {drive_file}")

        print(f"File uploaded successfully to Google Drive with ID: {drive_file.get('id')}")
        return drive_file.get('id')

    except Exception as e:
        print(f"Failed to upload file to Google Drive: {str(e)}")
        return None
