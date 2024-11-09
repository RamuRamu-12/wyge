import os

import PyPDF2
import openai
from docx import Document
from openai import OpenAI
import uuid
import io
from qdrant_client import QdrantClient
from qdrant_client.http import models
from dotenv import load_dotenv
load_dotenv()


openai.api_key = os.getenv("OPENAI_API_KEY")

# # Initialize Qdrant client
# qdrant_client = QdrantClient(
#     url="https://your-qdrant-url",
#     api_key="your-qdrant-api-key"
# )

# collection_name = "document_embeddings"  # Change as needed

def extract_text_from_pdf(pdf_file):
    """Extract text from a single PDF file"""
    text = ""
    reader = PyPDF2.PdfReader(pdf_file)
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    """Extract text from a single Word document"""
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

def extract_text_from_documents(files):
    """Extract text from multiple PDF and Word documents"""
    all_text = ""
    for file in files:
        file_copy = io.BytesIO(file.read())
        file_copy.name = file.name

        if file.name.lower().endswith('.pdf'):
            all_text += extract_text_from_pdf(file_copy)
        elif file.name.lower().endswith('.docx'):
            all_text += extract_text_from_docx(file_copy)
        else:
            raise ValueError(f"Unsupported file type: {file.name}")
    return all_text

def split_text_into_chunks(text, chunk_size):
    """Split text into chunks of specified size, trying to break at sentence boundaries"""
    chunks = []
    current_chunk = ""
    sentences = text.replace("\n", " ").split(". ")
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += sentence + ". "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

def text_to_embeddings_openai(chunks):
    """Generate embeddings using OpenAI"""
    embeddings = []
    for chunk in chunks:
        response = openai.embeddings.create(
            input=chunk,
            model="text-embedding-3-small",
            dimensions=384
        )
        embeddings.append(response.data[0].embedding)
    return chunks, embeddings

def store_in_qdrant(qdrant_client, collection_name, chunks, embeddings):
    """Store embeddings and chunks in Qdrant"""
    ids = [str(uuid.uuid4()) for _ in range(len(chunks))]
    points = [
        models.PointStruct(id=id, vector=embedding, payload={"text": chunk})
        for id, chunk, embedding in zip(ids, chunks, embeddings)
    ]
    
    qdrant_client.upsert(
        collection_name=collection_name,
        points=points
    )

def process_documents_and_store(qdrant_client, collection_name, files, chunk_size):
    """Handle document uploads, chunking, and embedding storage"""
    text = extract_text_from_documents(files)
    print(1)
    chunks = split_text_into_chunks(text, chunk_size)
    print(2)
    chunks, embeddings = text_to_embeddings_openai(chunks)
    print(3)
    store_in_qdrant(qdrant_client, collection_name, chunks, embeddings)
    print(4)

def query_to_embedding_openai(query_text):
    """Convert user query to embedding"""
    response = openai.embeddings.create(
        input=query_text,
        model="text-embedding-3-small",
        dimensions=384
    )
    return response.data[0].embedding

def query_qdrant(qdrant_client, collection_name, query_text, top_k=5):
    """Query Qdrant with the embedded query"""
    query_embedding = query_to_embedding_openai(query_text)
    results = qdrant_client.search(
        collection_name=collection_name,
        query_vector=query_embedding,
        limit=top_k,
        with_payload=True
    )
    return [(result.id, result.payload) for result in results]



# import PyPDF2
# from docx import Document
# from openai import OpenAI
# import uuid
# import io

# # Initialize OpenAI API key
# api_key = "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"
# openai_client = OpenAI(api_key=api_key)

# def extract_text_from_pdf(pdf_file):
#     """Extract text from a single PDF file"""
#     text = ""
#     reader = PyPDF2.PdfReader(pdf_file)
#     for page in reader.pages:
#         text += page.extract_text()
#     return text

# def extract_text_from_docx(docx_file):
#     """Extract text from a single Word document"""
#     doc = Document(docx_file)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 text += cell.text + " "
#             text += "\n"
#     return text

# def extract_text_from_documents(files):
#     """Extract text from multiple PDF and Word documents"""
#     all_text = ""
#     for file in files:
#         # Create a copy of the file in memory to avoid streamlit file handling issues
#         file_copy = io.BytesIO(file.getvalue())
#         file_copy.name = file.name

#         if file.name.lower().endswith('.pdf'):
#             all_text += extract_text_from_pdf(file_copy)
#         elif file.name.lower().endswith('.docx'):
#             all_text += extract_text_from_docx(file_copy)
#         else:
#             raise ValueError(f"Unsupported file type: {file.name}")
#     return all_text

# def split_text_into_chunks(text, chunk_size):
#     """Split text into chunks of specified size, trying to break at sentence boundaries"""
#     chunks = []
#     current_chunk = ""
#     sentences = text.replace("\n", " ").split(". ")
    
#     for sentence in sentences:
#         if len(current_chunk) + len(sentence) <= chunk_size:
#             current_chunk += sentence + ". "
#         else:
#             if current_chunk:
#                 chunks.append(current_chunk.strip())
#             current_chunk = sentence + ". "
    
#     if current_chunk:
#         chunks.append(current_chunk.strip())
    
#     return chunks

# def text_to_embeddings_openai(chunks):
#     """Generate embeddings using OpenAI"""
#     embeddings = []
#     for chunk in chunks:
#         response = openai_client.embeddings.create(
#             input=chunk,
#             model="text-embedding-3-small"
#         )
#         embeddings.append(response.data[0].embedding)
#     return chunks, embeddings

# def store_in_chromadb(chunks, embeddings, collection):
#     """Store embeddings and chunks in ChromaDB"""
#     ids = [str(uuid.uuid4()) for _ in range(len(chunks))]
#     collection.add(
#         documents=chunks,
#         embeddings=embeddings,
#         ids=ids 
#     )

# def process_documents_and_store(files, collection, chunk_size):
#     """Handle document uploads, chunking, and embedding storage"""
#     text = extract_text_from_documents(files)
#     chunks = split_text_into_chunks(text, chunk_size)
#     chunks, embeddings = text_to_embeddings_openai(chunks)
#     store_in_chromadb(chunks, embeddings, collection)

# def query_to_embedding_openai(query_text):
#     """Convert user query to embedding"""
#     response = openai_client.embeddings.create(
#         input=query_text,
#         model="text-embedding-3-small"
#     )
#     return response.data[0].embedding

# def query_chromadb(query_text, collection, top_k=5):
#     """Query ChromaDB with the embedded query"""
#     query_embedding = query_to_embedding_openai(query_text)
#     results = collection.query(
#         query_embeddings=[query_embedding],
#         n_results=top_k 
#     )
#     return results